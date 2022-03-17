# -*- coding: utf-8 -*-

from odoo import models, fields, api
from datetime import datetime
import calendar
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
import base64

DEFAULT_MONTH ={
	'january': 'Jan',
	'february': 'Feb',
	'march': 'Mar',
	'april': 'Apr',
	'may': 'May',
	'june': 'Jun',
	'july': 'Jul',
	'august': 'Aug',
	'september': 'Sep',
	'october': 'Oct',
	'november': 'Nov',
	'december': 'Dec'
}
class ApplicationDocx(models.TransientModel):
	"""Skill Wizard"""

	_name = "application.docx"
	_description = "Application Docx"

	name = fields.Char('Name')
	job_application = fields.Many2one('hr.applicant', 'Job Application')
	template_selection = fields.Selection([('naadee', 'Naadee Resume Template'), ('paskon', 'Paskon Resume Template')], 'Template', required=True, default='naadee')
	filename = fields.Char('File Name', readonly=True)
	filedata = fields.Binary('File', readonly=True)

	@api.model
	def default_get(self, fields):
		res = super(ApplicationDocx, self).default_get(fields)
		job_application = self.env['hr.applicant'].browse(self.env.context.get('active_ids'))
		res.update({
			'job_application': job_application.id
		})
		return res

	# @api.multi
	def appication_dox_report(self):
		active_ids = self.env.context.get('active_ids', []) or []
		job_application_id = self.job_application
		company_id = self.env.user.company_id
		document = Document()
		if self.template_selection == 'naadee':
			"""Template start from here."""
			logo_path = "/opt/odoo12/custom/solvithub_ras/static/description/company_logo.png"
			section = document.sections[0]
			sec_header = section.header
			header_tp = sec_header.add_paragraph()
			header_run = header_tp.add_run()
			header_run.add_picture(logo_path, width=Inches(1.0))
			if company_id.street:
				rml_header = "\t" + company_id.street
				header_run.add_text(rml_header)
			if company_id.city:
				rml_header = ", " + company_id.city
				header_run.add_text(rml_header)
			if company_id.state_id:
				rml_header = ", " + company_id.state_id.name
				header_run.add_text(rml_header)
			header_run.add_text("\n")
			header_run.add_text("\t\t\t\t\t\t")
			if company_id.phone:
				rml_header = "Office:" + company_id.phone + " "
				header_run.add_text(rml_header)
			rml_header = "Fax:" + company_id.phone
			header_run.add_text(rml_header)
			if company_id.website:
				rml_header = " " + company_id.website
				header_run.add_text(rml_header)
			font = header_run.font
			font.color.rgb = RGBColor(8, 80, 196)
			header_run.font.size = Pt(8)

			section1 = document.sections[0]
			sec_header1 = section1.header
			header_tp1 = sec_header1.add_paragraph()
			header_run1 = header_tp1
			header_run1.add_run("\t\t\t\t\t" + job_application_id.partner_name)
			# document.add_heading("\t\t\t\t\t" + job_application_id.partner_name, level=2)
			header_run1.add_run("\n_________________________________________________________________________________________________________")

			document.add_heading('SUMMARY', 2)
			if job_application_id.report_objective:
				paragraph = document.add_paragraph(job_application_id.report_objective)
			else:
				paragraph = document.add_paragraph("None")

			document.add_heading('EDUCATION', 2)
			if job_application_id.educational_details:
				for line in job_application_id.educational_details:
					table_edu = document.add_table(rows=len(line), cols=2)
					table_edu.columns[0].width = 3200400
					tx_cells = table_edu.rows[0].cells
					tb_cell_run = tx_cells[0]
					degree = ''
					if line.type_id:
						degree = line.type_id.name
					qualification = ''
					if line.qual_name:
						qualification = line.qual_name.name
					major = ''
					if line.qual_major:
						major = line.qual_major.name
					type = ''
					if line.school_name:
						major = line.school_name
					university = ''
					if line.school_university:
						university = line.school_university
					tb_cell_run.text = ("Degree: " + degree)
					tb_cell_run.text += ("\n")
					tb_cell_run.text += ("Qualification: " + qualification)
					tb_cell_run.text += ("\n")
					tb_cell_run.text += ("Qualification Major: " + major)
					tb_cell_run.text += ("\n")
					tb_cell_run.text += ("School Type: " + type)
					tb_cell_run.text += ("\n")
					tb_cell_run.text += ("School University: " + university)
					start_month = None
					start_year = None
					start_month_name = None
					end_month = None
					end_year = None
					end_month_name = None
					if line.start_date:
						start_month = datetime.strptime(str(line.start_date), '%Y-%m-%d').strftime('%m')
						start_year = datetime.strptime(str(line.start_date), '%Y-%m-%d').strftime('%Y')
						start_month_name = calendar.month_name[int(start_month)]
						start_month_name = start_month_name.lower()
						start_month_name = DEFAULT_MONTH[start_month_name]
					if line.end_date:
						end_month = datetime.strptime(str(line.end_date), '%Y-%m-%d').strftime('%m')
						end_year = datetime.strptime(str(line.end_date), '%Y-%m-%d').strftime('%Y')
						end_month_name = calendar.month_name[int(end_month)]
						end_month_name = end_month_name.lower()
						end_month_name = DEFAULT_MONTH[end_month_name]
					tb_cell_run_1 = tx_cells[1]
					tb_cell_run_1.text = ("\t\t")
					tb_cell_run_1.text += (start_month_name + ", " + start_year + " - " + end_month_name + ", " + end_year)

			else:
				paragraph = document.add_paragraph("None")

			document.add_heading('WORK EXPERIENCE', 2)
			if job_application_id.employment_details:
				for line in job_application_id.employment_details:
					table_emp = document.add_table(rows=len(line), cols=2)
					table_emp.columns[0].width = 3200400
					emp_cells = table_emp.rows[0].cells
					emp_cell_run = emp_cells[0]
					emp_cell_run.text = (line.name)
					emp_cell_run.text += ("\n")
					emp_cell_run.text += ("Present Position: " + line.position)
					emp_cell_run.text += ("\n")
					emp_cell_run.text += (line.description)
					start_month = None
					start_year = None
					start_month_name = None
					end_month = None
					end_year = None
					end_month_name = None
					if line.start_date:
						start_month = datetime.strptime(str(line.start_date), '%Y-%m-%d').strftime('%m')
						start_year = datetime.strptime(str(line.start_date), '%Y-%m-%d').strftime('%Y')
						start_month_name = calendar.month_name[int(start_month)]
						start_month_name = start_month_name.lower()
						start_month_name = DEFAULT_MONTH[start_month_name]
					if line.end_date:
						end_month = datetime.strptime(str(line.end_date), '%Y-%m-%d').strftime('%m')
						end_year = datetime.strptime(str(line.end_date), '%Y-%m-%d').strftime('%Y')
						end_month_name = calendar.month_name[int(end_month)]
						end_month_name = end_month_name.lower()
						end_month_name = DEFAULT_MONTH[end_month_name]
					emp_cell_run_1 = emp_cells[1]
					emp_cell_run_1.text = ("\t\t")
					emp_cell_run_1.text += (start_month_name + ", " + start_year + " - " + end_month_name + ", " + end_year)
			else:
				paragraph = document.add_paragraph("None")

			document.add_heading('CERTIFICATIONS', 2)
			if job_application_id.certifications:
				for line in job_application_id.certifications:
					document.add_paragraph(line.name, style = 'List Bullet')

		elif self.template_selection == 'paskon':
			pass

		# document.add_page_break()
		document.save("/opt/Test_word.docx")
		file = open("/opt/Test_word.docx", "rb")
		out = file.read()
		self.filedata = base64.encodestring(out)
		self.filename = job_application_id.partner_name + "_" + self.template_selection + ".docx"
		file.close()
		return {
			'name': 'Application Document Report',
			'res_model': 'application.docx',
			'type': 'ir.actions.act_window',
			'view_type': 'form',
			'view_mode': 'form',
			'target': 'new',
			'nodestroy': True,
			'res_id': self.id,
		}